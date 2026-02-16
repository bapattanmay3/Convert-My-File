# Keep ONLY built-in imports at top
import os
import re
import time
import csv
# All other imports moved inside functions

# ===== GLOBAL LANGUAGES CONSTANT =====
LANGUAGES = {
    'af': 'Afrikaans', 'sq': 'Albanian', 'am': 'Amharic', 'ar': 'Arabic', 'hy': 'Armenian',
    'az': 'Azerbaijani', 'eu': 'Basque', 'be': 'Belarusian', 'bn': 'Bengali', 'bs': 'Bosnian',
    'bg': 'Bulgarian', 'ca': 'Catalan', 'ceb': 'Cebuano', 'ny': 'Chichewa', 'zh-cn': 'Chinese (Simplified)',
    'zh-tw': 'Chinese (Traditional)', 'co': 'Corsican', 'hr': 'Croatian', 'cs': 'Czech', 'da': 'Danish',
    'nl': 'Dutch', 'en': 'English', 'eo': 'Esperanto', 'et': 'Estonian', 'tl': 'Filipino',
    'fi': 'Finnish', 'fr': 'French', 'fy': 'Frisian', 'gl': 'Galician', 'ka': 'Georgian',
    'de': 'German', 'el': 'Greek', 'gu': 'Gujarati', 'ht': 'Haitian Creole', 'ha': 'Hausa',
    'haw': 'Hawaiian', 'iw': 'Hebrew', 'hi': 'Hindi', 'hmn': 'Hmong', 'hu': 'Hungarian',
    'is': 'Icelandic', 'ig': 'Igbo', 'id': 'Indonesian', 'ga': 'Irish', 'it': 'Italian',
    'ja': 'Japanese', 'jw': 'Javanese', 'kn': 'Kannada', 'kk': 'Kazakh', 'km': 'Khmer',
    'ko': 'Korean', 'ku': 'Kurdish (Kurmanji)', 'ky': 'Kyrgyz', 'lo': 'Lao', 'la': 'Latin',
    'lv': 'Latvian', 'lt': 'Lithuanian', 'lb': 'Luxembourgish', 'mk': 'Macedonian', 'mg': 'Malagasy',
    'ms': 'Malay', 'ml': 'Malayalam', 'mt': 'Maltese', 'mi': 'Maori', 'mr': 'Marathi',
    'mn': 'Mongolian', 'my': 'Myanmar (Burmese)', 'ne': 'Nepali', 'no': 'Norwegian', 'ps': 'Pashto',
    'fa': 'Persian', 'pl': 'Polish', 'pt': 'Portuguese', 'pa': 'Punjabi', 'ro': 'Romanian',
    'ru': 'Russian', 'sm': 'Samoan', 'gd': 'Scots Gaelic', 'sr': 'Serbian', 'st': 'Sesotho',
    'sn': 'Shona', 'sd': 'Sindhi', 'si': 'Sinhala', 'sk': 'Slovak', 'sl': 'Slovenian',
    'so': 'Somali', 'es': 'Spanish', 'su': 'Sundanese', 'sw': 'Swahili', 'sv': 'Swedish',
    'tg': 'Tajik', 'ta': 'Tamil', 'te': 'Telugu', 'th': 'Thai', 'tr': 'Turkish',
    'uk': 'Ukrainian', 'ur': 'Urdu', 'uz': 'Uzbek', 'vi': 'Vietnamese', 'cy': 'Welsh',
    'xh': 'Xhosa', 'yi': 'Yiddish', 'yo': 'Yoruba', 'zu': 'Zulu'
}

# ===== ENTITY PRESERVATION PATTERNS =====
PRESERVE_PATTERNS = [
    r'^[\d\s\+\-\(\)]+$',                    # Pure numbers
    r'^\d+$',                                 # Integers
    r'^\d+\.\d+$',                           # Decimals
    r'^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}$',    # Dates
    r'^[\w\.-]+@[\w\.-]+\.\w+$',             # Emails
    r'^https?://[^\s]+$',                    # URLs
    r'^[A-Z0-9]{5,}$',                       # Codes/IDs
    r'^[A-Z]{2,}\d+$',                       # Alphanumeric codes
    r'^\$\s*\d+(\.\d{2})?$',                 # Prices
    r'^\d+(\.\d{1,2})?%$',                   # Percentages
]

def clean_text(text):
    """Normalize text by removing redundant whitespace and fixing spaced-out PDF text"""
    if not text: return ""
    # Replace null bytes
    text = text.replace('\x00', '')
    
    # Handle spaced-out text (e.g. "P r o t e c t e d") which confuses translators
    # Heuristic: If we see a pattern of [Letter][Space][Letter][Space]
    if len(text) > 10:
        # Use regex to find sequences of single letters separated by spaces
        # e.g., "H e l l o" -> "Hello"
        spaced_pattern = r'(?:(?<=\s)|(?<=^))([a-zA-Z0-9])\s+(?=([a-zA-Z0-9])(?:\s|$))'
        # Group 1 is the character. We replace the character + space with just the character.
        text = re.sub(spaced_pattern, r'\1', text)
    
    # Final normalization
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def should_preserve(text):
    """Check if text should be kept as-is (not translated)"""
    if not isinstance(text, str):
        return True
    text = text.strip()
    if not text or len(text) < 2:
        return True
    for pattern in PRESERVE_PATTERNS:
        if re.match(pattern, text):
            return True
    # Check if mostly numbers
    letters = sum(c.isalpha() for c in text)
    numbers = sum(c.isdigit() for c in text)
    if numbers > letters and numbers > 3:
        return True
    return False

# ===== TRANSLATE FUNCTION USING TRANSLATORS LIBRARY =====
def translate_text(text, target_lang, source_lang='auto', max_retries=3):
    """Safe translation with UTF-8 preservation"""
    from deep_translator import GoogleTranslator  # ✅ ADD THIS
    if not text or should_preserve(text):
        return text
    
    # Ensure text is properly encoded
    if isinstance(text, bytes):
        text = text.decode('utf-8', errors='ignore')
    
    # Normalize text (keeping my de-spacing logic too as it proved useful)
    text = clean_text(text)
    
    # Special handling for language codes (Google/Deep-Translator variants)
    target = target_lang.lower()
    if '-' in target:
        # e.g., zh-cn -> zh-CN, pt-br -> pt-BR
        parts = target.split('-')
        target = f"{parts[0]}-{parts[1].upper()}"
    elif target in ['iw', 'he']:
        target = 'he' # Hebrew
    elif target == 'ms':
        target = 'ms' # Malay is fine
    
    chunk_size = 4000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    translated_chunks = []
    
    for chunk in chunks:
        for attempt in range(max_retries):
            try:
                translator = GoogleTranslator(source=source_lang, target=target)
                result = translator.translate(chunk)
                if result:
                    # Force UTF-8
                    if isinstance(result, str):
                        result = result.encode('utf-8').decode('utf-8')
                    translated_chunks.append(result)
                    break
                time.sleep(1)
            except Exception as e:
                try:
                    print(f"Translation attempt {attempt+1} failed: {e}")
                except UnicodeEncodeError:
                    pass
                if attempt == max_retries - 1:
                    # Fallback with UTF-8 encoding
                    fallback = chunk.encode('utf-8').decode('utf-8')
                    translated_chunks.append(fallback)
                time.sleep(2)
    
    return ' '.join(translated_chunks)

def is_valid_translation(translated_text, target_lang, original_text=None):
    """
    Check if the translation is valid for the target language.
    1. For languages with distinct scripts, checks Unicode ranges.
    2. For Latin-based languages, checks if it's different from original (if provided).
    """
    if not translated_text:
        return False
    
    target = target_lang.split('-')[0].lower()
    
    # Script-based validation mapping
    script_ranges = {
        'hi': ('\u0900', '\u097F'),  # Devanagari (Hindi, Marathi, etc.)
        'bn': ('\u0980', '\u09FF'),  # Bengali
        'ar': ('\u0600', '\u06FF'),  # Arabic
        'fa': ('\u0600', '\u06FF'),  # Persian
        'ur': ('\u0600', '\u06FF'),  # Urdu
        'ru': ('\u0400', '\u04FF'),  # Cyrillic (Russian, etc.)
        'uk': ('\u0400', '\u04FF'),  # Ukrainian
        'be': ('\u0400', '\u04FF'),  # Belarusian
        'el': ('\u0370', '\u03FF'),  # Greek
        'iw': ('\u0590', '\u05FF'),  # Hebrew
        'he': ('\u0590', '\u05FF'),  # Hebrew (alternative code)
        'ja': ('\u3040', '\u9FFF'),  # Japanese (Hiragana/Katakana/Kanji)
        'zh': ('\u4E00', '\u9FFF'),  # Chinese (CJK Unified Ideographs)
        'ko': ('\uAC00', '\uD7AF'),  # Korean (Hangul)
        'ta': ('\u0B80', '\u0BFF'),  # Tamil
        'te': ('\u0C00', '\u0C7F'),  # Telugu
        'kn': ('\u0C80', '\u0CFF'),  # Kannada
        'ml': ('\u0D00', '\u0D7F'),  # Malayalam
        'gu': ('\u0A80', '\u0AFF'),  # Gujarati
        'pa': ('\u0A00', '\u0A7F'),  # Punjabi
        'th': ('\u0E00', '\u0E7F'),  # Thai
    }
    
    if target in script_ranges:
        low, high = script_ranges[target]
        chars_in_range = sum(1 for c in translated_text if low <= c <= high)
        return chars_in_range > len(translated_text.strip()) * 0.2  # 20% threshold
    
    # For Latin-based or others, check if it's different from original
    if original_text:
        # Simple similarity check: if they are identical, it's likely failed
        if translated_text.strip() == original_text.strip() and len(original_text) > 10:
            return False
            
    return True

# ===== PDF TRANSLATOR (Structural Bridge) =====
def translate_pdf(input_path, output_path, target_lang, source_lang='auto'):
    """
    Translate PDF while preserving tables and layout.
    Bridge: PDF -> DOCX (Structural) -> Translate -> PDF
    """
    from pdf2docx import Converter
    import os
    
    unique_id = os.path.basename(input_path).split('_')[0]
    temp_docx = input_path + ".structural.docx"
    temp_translated_docx = input_path + ".translated.docx"
    
    try:
        # 1. Convert PDF to structural DOCX (preserves tables/layout)
        cv = Converter(input_path)
        cv.convert(temp_docx, start=0, end=None)
        cv.close()
        
        # 2. Translate the structural DOCX
        success, msg = translate_docx(temp_docx, temp_translated_docx, target_lang, source_lang)
        if not success:
            return False, f"Structural translation failed: {msg}"
            
        # 3. Convert translated DOCX back to PDF (using weasyprint bridge for script support)
        # Handle conversion via pypandoc to HTML then to PDF for best language script rendering
        import pypandoc
        temp_html = input_path + ".translated.html"
        pypandoc.convert_file(temp_translated_docx, 'html', outputfile=temp_html)
        
        from weasyprint import HTML
        HTML(temp_html).write_pdf(output_path)
        
        # Cleanup
        for tmp in [temp_docx, temp_translated_docx, temp_html]:
            if os.path.exists(tmp): os.remove(tmp)
            
        return True, "PDF structural translation completed successfully"
        
    except Exception as e:
        print(f"Structural PDF translation error: {e}")
        # Fallback to Text if structural bridge fails
        output_txt = output_path.replace('.pdf', '.txt')
        return translate_pdf_to_text_fallback(input_path, output_txt, target_lang, source_lang)

def translate_pdf_to_text_fallback(input_path, output_path, target_lang, source_lang='auto'):
    """Fallback plain-text translation if structural bridge fails"""
    import pdfplumber
    try:
        text_content = []
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(translate_text(text, target_lang, source_lang))
        
        with open(output_path, 'w', encoding='utf-8-sig') as f:
            f.write('\n\n'.join(text_content))
        return True, "PDF translated to text (Structure preservation failed)"
    except Exception as e:
        return False, str(e)

def translate_docx(input_path, output_path, target_lang, source_lang='auto'):
    """Translate Word documents with batching to prevent timeouts/502s"""
    try:
        from docx import Document
        doc = Document(input_path)
        
        def process_batch(paragraphs, current_batch_text, batch_indices):
            if not current_batch_text:
                return
            
            # Combine batch with a unique separator that's unlikely to be in text
            separator = " [[[DSEP]]] "
            combined_text = separator.join(current_batch_text)
            
            translated_combined = translate_text(combined_text, target_lang, source_lang)
            translated_parts = translated_combined.split(separator)
            
            # If split count matches, apply translations
            if len(translated_parts) == len(batch_indices):
                for idx, trans in zip(batch_indices, translated_parts):
                    paragraphs[idx].text = trans.strip()
            else:
                # Fallback: translate individually if batching fails
                for idx in batch_indices:
                    try:
                        paragraphs[idx].text = translate_text(paragraphs[idx].text, target_lang, source_lang).strip()
                    except: pass
            
            # Anti-throttle breather
            time.sleep(0.5)

        # 1. Translate Main Paragraphs in Batches
        all_paras = list(doc.paragraphs)
        current_batch = []
        current_indices = []
        current_length = 0
        
        for i, para in enumerate(all_paras):
            text = para.text.strip()
            if text and not should_preserve(text):
                current_batch.append(text)
                current_indices.append(i)
                current_length += len(text)
                
                # Batch limit: ~3000 chars or 20 paragraphs
                if current_length > 3000 or len(current_batch) >= 20:
                    process_batch(all_paras, current_batch, current_indices)
                    current_batch, current_indices, current_length = [], [], 0
        
        # Process remaining
        process_batch(all_paras, current_batch, current_indices)
        
        # 2. Translate Tables (CRITICAL: Extract and translate while keeping structure)
        for table in doc.tables:
            for row in table.rows:
                table_batch = []
                cell_refs = []
                for cell in row.cells:
                    # Collect cell text for batching per row to stay efficient but safe
                    cell_text = cell.text.strip()
                    if cell_text and not should_preserve(cell_text):
                        table_batch.append(cell_text)
                        cell_refs.append(cell)
                
                if table_batch:
                    separator = " [[[TSEP]]] "
                    combined = separator.join(table_batch)
                    translated_combined = translate_text(combined, target_lang, source_lang)
                    translated_parts = translated_combined.split(separator)
                    
                    if len(translated_parts) == len(cell_refs):
                        for cell, part in zip(cell_refs, translated_parts):
                            cell.text = part.strip()
                    else:
                        for cell in cell_refs:
                            try:
                                cell.text = translate_text(cell.text, target_lang, source_lang).strip()
                            except: pass
                    
                    # Anti-throttle breather
                    time.sleep(0.5)
        
        doc.save(output_path)
        return True, "Word document translation completed"
    except Exception as e:
        print(f"DOCX translation error: {e}")
        import traceback
        traceback.print_exc()
        return False, str(e)

def translate_excel(input_path, output_path, target_lang, source_lang='auto'):
    """Translate Excel files with batching and legacy .xls support"""
    import os
    file_ext = os.path.splitext(input_path)[1].lower()
    temp_xlsx = None
    
    try:
        import openpyxl
        import pandas as pd
        
        # Bridge legacy .xls to .xlsx
        if file_ext == '.xls':
            temp_xlsx = input_path + ".bridge.xlsx"
            # Read all sheets from .xls and save to .xlsx
            with pd.ExcelWriter(temp_xlsx, engine='openpyxl') as writer:
                # Use xlrd engine for old .xls
                xls_data = pd.read_excel(input_path, sheet_name=None, engine='xlrd')
                for sheet_name, df in xls_data.items():
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            current_input = temp_xlsx
        else:
            current_input = input_path

        wb = openpyxl.load_workbook(current_input)
        
        # 1. Collect all translatable cells
        cells_to_translate = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    # Only translate strings, ignore formulas and None
                    if cell.value and isinstance(cell.value, str) and not cell.data_type == 'f':
                        text = cell.value.strip()
                        if text and not should_preserve(text):
                            cells_to_translate.append((cell, text))
        
        # 2. Batch process
        if cells_to_translate:
            current_batch_texts = []
            current_batch_cells = []
            current_len = 0
            
            def process_excel_batch(texts, cells):
                # Unique separator that is unlikely to be mangled by translation
                sep = " [[[XSEP]]] "
                combined = sep.join(texts)
                translated = translate_text(combined, target_lang, source_lang)
                parts = translated.split(sep)
                
                if len(parts) == len(cells):
                    for cell, trans in zip(cells, parts):
                        cell.value = trans.strip()
                else:
                    # Individual fallback if separator fails
                    for cell in cells:
                        try:
                            cell.value = translate_text(str(cell.value), target_lang, source_lang)
                        except: pass
                
                # Small breather to avoid API rate limiting/timeouts
                time.sleep(0.5)

            for cell, text in cells_to_translate:
                current_batch_texts.append(text)
                current_batch_cells.append(cell)
                current_len += len(text)
                
                # Smaller batches for stability
                if current_len > 2000 or len(current_batch_texts) >= 20:
                    process_excel_batch(current_batch_texts, current_batch_cells)
                    current_batch_texts, current_batch_cells, current_len = [], [], 0
            
            # Final batch
            if current_batch_texts:
                process_excel_batch(current_batch_texts, current_batch_cells)
        
        wb.save(output_path)
        
        # Cleanup bridge
        if temp_xlsx and os.path.exists(temp_xlsx):
            os.remove(temp_xlsx)
            
        return True, "Excel translation completed"
    except Exception as e:
        if temp_xlsx and os.path.exists(temp_xlsx):
            os.remove(temp_xlsx)
        print(f"Excel translation error: {e}")
        return False, f"Excel error: {str(e)}"

def translate_csv(input_path, output_path, target_lang, source_lang='auto'):
    """Translate CSV files with row-based batching"""
    try:
        with open(input_path, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            rows = list(reader)
        
        if not rows:
            return True, "Empty CSV"

        translated_rows = []
        for row in rows:
            # Batch translate the entire row
            translatable_indices = []
            row_texts = []
            
            for i, cell in enumerate(row):
                if cell and not should_preserve(cell):
                    translatable_indices.append(i)
                    row_texts.append(cell)
            
            if row_texts:
                sep = " ||| "
                translated = translate_text(sep.join(row_texts), target_lang, source_lang)
                parts = translated.split(sep)
                
                new_row = list(row)
                if len(parts) == len(translatable_indices):
                    for idx, trans in zip(translatable_indices, parts):
                        new_row[idx] = trans.strip()
                else:
                    # Fallback
                    for idx in translatable_indices:
                        new_row[idx] = translate_text(row[idx], target_lang, source_lang)
                translated_rows.append(new_row)
            else:
                translated_rows.append(row)
                
        with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            writer.writerows(translated_rows)
            
        return True, "CSV translation completed"
    except Exception as e:
        return False, str(e)

# ===== TEXT FILE TRANSLATOR =====
def translate_text_file(input_path, output_path, target_lang, source_lang='auto'):
    """Translate plain text files"""
    try:
        with open(input_path, 'r', encoding='utf-8-sig') as f:
            content = f.read()
        
        translated = translate_text(content, target_lang, source_lang)
        
        # Force UTF-8 with BOM
        with open(output_path, 'w', encoding='utf-8-sig') as f:
            f.write(translated)
        return True, "Text translation completed"
    except Exception as e:
        try:
            print(f"Text translation error: {e}")
        except UnicodeEncodeError:
            pass
        return False, str(e)

# ===== MAIN DISPATCHER FUNCTION =====
def translate_document(input_path, output_path, target_lang, source_lang='auto', file_ext=None):
    """Main dispatcher function - supports all formats"""
    if file_ext is None:
        file_ext = os.path.splitext(input_path)[1].lower()
    
    # Map file extensions to translator functions
    translators = {
        '.pdf': translate_pdf,
        '.docx': translate_docx,
        '.doc': translate_docx,  # Same as docx
        '.xlsx': translate_excel,
        '.xls': translate_excel,   # Same as xlsx
        '.csv': translate_csv,      
        '.txt': translate_text_file 
    }
    
    translator = translators.get(file_ext)
    
    # Special case: Legacy .doc bridge
    if file_ext == '.doc':
        try:
            import pypandoc
            temp_docx = input_path + ".bridge.docx"
            # Ensure pypandoc uses the system pandoc
            pypandoc.convert_file(input_path, 'docx', outputfile=temp_docx)
            
            # Translate the bridge file
            success, message = translate_docx(temp_docx, output_path, target_lang, source_lang)
            
            # Cleanup bridge
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            return success, message
        except Exception as de:
            print(f"Legacy .doc conversion failed: {de}")
            return False, f"Legacy .doc support requires pandoc: {str(de)}"

    if translator:
        return translator(input_path, output_path, target_lang, source_lang)
    else:
        return False, f"Unsupported file type: {file_ext}"

# Backward compatibility (some files might still call translate_file)
def translate_file(input_path, output_path, target_lang, source_lang='auto'):
    return translate_document(input_path, output_path, target_lang, source_lang)