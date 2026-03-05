# Standard library imports
import os
import uuid
import urllib.parse

# Flask related imports
from flask import Flask, render_template, request, jsonify, send_from_directory, flash, redirect, url_for, send_file, session
from functools import wraps
from werkzeug.utils import secure_filename

# Note: heavy imports (translator_engine, converter_universal, etc.) 
# are moved inside routes or after app initialization.

import time
import threading
import json
import requests
from datetime import datetime, timedelta

app = Flask(__name__)
app.config['SECRET_KEY'] = 'royal-enfield-racing-green-2026' # Change this for production
app.config['ADMIN_USER'] = 'bapattanmay'
app.config['ADMIN_PASS'] = 'qazwsxedcrfv@123'
app.config['SITE_NAME'] = 'Convert My File'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ANALYTICS_FILE'] = 'analytics.json'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Authentication Decorator
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# --- Admin Routes ---

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if username == app.config['ADMIN_USER'] and password == app.config['ADMIN_PASS']:
            session['admin_logged_in'] = True
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid username or password', 'error')
            
    return render_template('login.html', site_name=app.config['SITE_NAME'])

@app.route('/logout')
def logout():
    session.pop('admin_logged_in', None)
    return redirect(url_for('login'))

# Initialize Analytics File
if not os.path.exists(app.config['ANALYTICS_FILE']):
    with open(app.config['ANALYTICS_FILE'], 'w') as f:
        json.dump({"visitors": [], "usage_count": 0}, f)

def log_visit():
    """Log a unique visitor visit with location data"""
    try:
        # Get public IP (Simple heuristic for demo/local testing)
        ip = request.headers.get('X-Forwarded-For', request.remote_addr)
        
        # Load current analytics
        with open(app.config['ANALYTICS_FILE'], 'r') as f:
            data = json.load(f)
        
        # Check if already logged (primitive session tracking)
        # In a real app, we'd use cookies or a database
        visitor_ips = [v.get('ip') for v in data['visitors']]
        if ip not in visitor_ips:
            # Get location via API (Free tier)
            location = "Unknown"
            try:
                # Using ip-api.com (no key needed for bulk/simple calls)
                res = requests.get(f"http://ip-api.com/json/{ip}", timeout=2).json()
                if res.get('status') == 'success':
                    location = f"{res.get('city')}, {res.get('country')}"
            except: pass
            
            data['visitors'].append({
                "ip": ip,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "location": location
            })
            
            with open(app.config['ANALYTICS_FILE'], 'w') as f:
                json.dump(data, f)
    except Exception as e:
        print(f"Tracking error: {e}")

def log_usage():
    """Increment the usage counter for successful service completion"""
    try:
        with open(app.config['ANALYTICS_FILE'], 'r') as f:
            data = json.load(f)
        data['usage_count'] = data.get('usage_count', 0) + 1
        with open(app.config['ANALYTICS_FILE'], 'w') as f:
            json.dump(data, f)
    except Exception as e:
        print(f"Usage logging error: {e}")

def cleanup_old_files():
    """Background task to delete files older than 5 minutes"""
    while True:
        try:
            now = time.time()
            # 300 seconds = 5 minutes
            threshold = now - 300
            
            for f in os.listdir(app.config['UPLOAD_FOLDER']):
                f_path = os.path.join(app.config['UPLOAD_FOLDER'], f)
                if os.path.isfile(f_path):
                    # Check file modification time
                    if os.path.getmtime(f_path) < threshold:
                        try:
                            os.remove(f_path)
                            print(f"Auto-deleted expired file: {f}")
                        except Exception as e:
                            print(f"Error deleting file {f}: {e}")
        except Exception as e:
            print(f"Cleanup thread error: {e}")
            
        time.sleep(60)  # Check every minute

# Start cleanup thread
cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()

@app.route('/robots.txt')
def robots():
    return send_from_directory('static', 'robots.txt')

@app.route('/sitemap.xml')
def sitemap():
    return send_from_directory('static', 'sitemap.xml')

@app.route('/admin')
@admin_required
def admin_dashboard():
    """Display analytics dashboard"""
    try:
        with open(app.config['ANALYTICS_FILE'], 'r') as f:
            data = json.load(f)
        
        visitors = data.get('visitors', [])
        total_visitors = len(visitors)
        usage_count = data.get('usage_count', 0)
        
        # Calculate Conversion Rate
        conversion_rate = 0
        if total_visitors > 0:
            conversion_rate = round((usage_count / total_visitors) * 100, 1)
            
        # Group by location
        locations = {}
        # Group by date
        daily_stats = {}
        
        for v in visitors:
            # Location stats
            loc = v.get('location', 'Unknown')
            locations[loc] = locations.get(loc, 0) + 1
            
            # Daily stats
            try:
                # v['timestamp'] is like "2026-03-05 15:45:12"
                date_str = v['timestamp'].split(' ')[0]
                daily_stats[date_str] = daily_stats.get(date_str, 0) + 1
            except: pass
            
        # Sort locations by count
        sorted_locations = dict(sorted(locations.items(), key=lambda item: item[1], reverse=True))
        # Sort daily stats by date (newest first)
        sorted_daily = dict(sorted(daily_stats.items(), key=lambda item: item[0], reverse=True))
        
        return render_template('admin.html', 
                             total_visitors=total_visitors,
                             usage_count=usage_count,
                             conversion_rate=conversion_rate,
                             locations=sorted_locations,
                             daily_stats=sorted_daily,
                             recent_visitors=visitors[-10:][::-1], # Last 10
                             site_name=app.config['SITE_NAME'])
    except Exception as e:
        return f"Dashboard error: {e}", 500

@app.route('/')
def home():
    log_visit()
    return render_template('index.html', site_name=app.config['SITE_NAME'])

@app.route('/favicon.ico')
def favicon():
    return '', 204

@app.route('/converter')
def converter():
    return render_template('converter.html', site_name=app.config['SITE_NAME'])

@app.route('/translator')
def translator():
    from translator_engine import LANGUAGES
    return render_template('translator.html', site_name=app.config['SITE_NAME'], languages=LANGUAGES)

@app.route('/merger')
def merger():
    return render_template('merger.html', site_name=app.config['SITE_NAME'])

@app.route('/merge', methods=['POST'])
def merge_pdfs():
    """Handle PDF and Image merging into a single PDF"""
    if 'files' not in request.files:
        flash('No files uploaded')
        return redirect(url_for('merger'))
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        flash('No files selected')
        return redirect(url_for('merger'))
    
    from PyPDF2 import PdfMerger
    from PIL import Image
    import io
    
    merger = PdfMerger()
    unique_id = str(uuid.uuid4())
    temp_files = []
    
    try:
        for file in files:
            filename = secure_filename(file.filename)
            ext = os.path.splitext(filename)[1].lower()
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"tmp_{unique_id}_{filename}")
            file.save(temp_path)
            temp_files.append(temp_path)
            
            if ext in ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff']:
                # Bridge: Convert image to temporary PDF for merging
                img_pdf_path = temp_path + ".bridge.pdf"
                img = Image.open(temp_path)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                img.save(img_pdf_path, "PDF")
                merger.append(img_pdf_path)
                temp_files.append(img_pdf_path)
            else:
                merger.append(temp_path)
        
        output_filename = f"merged_{unique_id}.pdf"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        with open(output_path, 'wb') as f:
            merger.write(f)
        
        merger.close()
        
        # Comprehensive Cleanup
        for tmp in temp_files:
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except:
                pass
                
        return send_from_directory(app.config['UPLOAD_FOLDER'], output_filename, as_attachment=True)
    
    except Exception as e:
        # Cleanup on failure
        for tmp in temp_files:
            try:
                if os.path.exists(tmp): os.remove(tmp)
            except: pass
        flash(f'Merging failed: {str(e)}')
        return redirect(url_for('merger'))

@app.route('/compressor')
def compressor():
    return render_template('compressor.html', site_name=app.config['SITE_NAME'])

@app.route('/compress', methods=['POST'])
def compress():
    """Handle file compression with target size (PDF and Images)"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    try:
        target_val = float(request.form.get('target_size', 0))
        unit = request.form.get('unit', 'KB').upper()
        
        # Target size in bytes
        target_bytes = target_val * 1024 if unit == 'KB' else target_val * 1024 * 1024
        
        # Enforcement: 5KB to 50MB
        min_bytes = 5 * 1024
        max_bytes = 50 * 1024 * 1024
        
        if target_bytes < min_bytes or target_bytes > max_bytes:
            return jsonify({'success': False, 'error': 'We can process 5KB-50MB files'}), 400

        # Save uploaded file
        filename = secure_filename(file.filename)
        unique_id = str(uuid.uuid4())
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(input_path)
        
        # Get extension
        ext = os.path.splitext(filename)[1].lower().replace('.', '')
        output_filename = f"processed_{unique_id}_{filename}"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        success = False
        message = ""

        if ext in ['jpg', 'jpeg', 'png', 'webp']:
            from PIL import Image
            import io
            
            orig_img = Image.open(input_path)
            orig_format = 'JPEG' if ext in ['jpg', 'jpeg'] else ext.upper()
            
            # --- Lever 1: Quality Tuning ---
            low_q, high_q = 1, 100
            best_q = 75
            
            for _ in range(7):
                mid_q = (low_q + high_q) // 2
                buf = io.BytesIO()
                
                # Parameters for format
                params = {'format': orig_format, 'optimize': True}
                if orig_format in ['JPEG', 'WEBP']: params['quality'] = mid_q
                elif orig_format == 'PNG': params['compress_level'] = mid_q // 11
                
                temp_img = orig_img
                if temp_img.mode != 'RGB' and orig_format == 'JPEG':
                    temp_img = temp_img.convert('RGB')
                    
                temp_img.save(buf, **params)
                if buf.tell() <= target_bytes:
                    best_q = mid_q
                    low_q = mid_q + 1
                else:
                    high_q = mid_q - 1

            # Check if Lever 1 was enough
            final_buf = io.BytesIO()
            test_params = {'format': orig_format, 'optimize': True}
            if orig_format in ['JPEG', 'WEBP']: test_params['quality'] = best_q
            elif orig_format == 'PNG': test_params['compress_level'] = best_q // 11
            
            temp_img = orig_img
            if temp_img.mode != 'RGB' and orig_format == 'JPEG':
                temp_img = temp_img.convert('RGB')
            temp_img.save(final_buf, **test_params)
            
            current_size = final_buf.tell()
            
            # --- Lever 2: Precision Convergence Loop (v3.0) ---
            # We iteratively recalibrate dimensions based on actual size feedback
            current_scale = 1.0
            tolerance = 0.05 # 5% error margin
            working_img = orig_img
            
            for attempt in range(5):
                # Calculate dimensions for this pass
                temp_w = max(1, int(orig_img.width * current_scale))
                temp_h = max(1, int(orig_img.height * current_scale))
                
                # Resize from ORIGINAL to avoid compound blur
                working_img = orig_img.resize((temp_w, temp_h), Image.Resampling.LANCZOS)
                if working_img.mode != 'RGB' and orig_format == 'JPEG':
                    working_img = working_img.convert('RGB')
                
                # Check output size
                buf = io.BytesIO()
                working_img.save(buf, **test_params)
                current_size = buf.tell()
                
                # Are we close enough?
                if abs(current_size - target_bytes) / target_bytes <= tolerance:
                    break
                
                # Recalculate scale based on actual size result
                # Ratio is sqrt(Target/Actual) because size is area-proportional
                size_ratio = target_bytes / current_size
                # Dampen the adjustment to avoid oscillations
                scale_adjustment = (size_ratio ** 0.5) * 0.95 if size_ratio < 1 else (size_ratio ** 0.5)
                current_scale *= scale_adjustment
                
                # Safety boundaries (0.01x to 8.0x original)
                current_scale = max(0.01, min(8.0, current_scale))

            # Final Save
            working_img.save(output_path, **test_params)
            success, message = True, f"Precision calibrated ({unit}) after {attempt+1} passes"
            
        elif ext == 'pdf':
            # PDF compression is less granular with PyPDF2
            from PyPDF2 import PdfReader, PdfWriter
            reader = PdfReader(input_path)
            writer = PdfWriter()
            for page in reader.pages:
                page.compress_content_streams()
                writer.add_page(page)
            
            with open(output_path, 'wb') as f:
                writer.write(f)
            success, message = True, "PDF processed with standard compression"
        else:
            success, message = False, f"Format {ext} not supported for target-size processing"

        if success and os.path.exists(output_path):
            # Usage tracked
            log_usage()
            return jsonify({
                'success': True,
                'message': message,
                'download_url': f'/download/{output_filename}',
                'filename': output_filename
            })
        else:
            return jsonify({'success': False, 'error': message or "Processing failed"}), 500
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/convert', methods=['POST'])
def convert_file_universal():
    """Universal file converter - handles all formats"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    target_format = request.form.get('format', '').lower()
    
    # Save uploaded file
    filename = secure_filename(file.filename)
    unique_id = str(uuid.uuid4())
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
    file.save(input_path)
    
    # Get file extension
    source_format = os.path.splitext(filename)[1].lower().replace('.', '')
    
    # Generate output filename
    base_name = os.path.splitext(filename)[0]
    output_filename = f"converted_{unique_id}_{base_name}.{target_format}"
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
    
    # Import converter
    import converter_universal as cv
    
    # Perform conversion
    try:
        if source_format in cv.FILE_CONVERSIONS and target_format in cv.FILE_CONVERSIONS[source_format]:
            conversion_func = cv.FILE_CONVERSIONS[source_format][target_format]
            success, message = conversion_func(input_path, output_path)
        else:
            success, message = False, f"Unsupported conversion from {source_format} to {target_format}"
    except Exception as e:
        success, message = False, str(e)
    
    if success and os.path.exists(output_path):
        return jsonify({
            'success': True,
            'message': message,
            'download_url': f'/download/{output_filename}',
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'error': message}), 500

@app.route('/convert-image', methods=['POST'])
def convert_image():
    """Handle image conversions (JPG, PNG, WebP, PDF)"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    target_format = request.form.get('format', 'jpg')
    quality = int(request.form.get('quality', 90))
    
    # Save uploaded file
    filename = secure_filename(file.filename)
    unique_id = str(uuid.uuid4())
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
    file.save(input_path)
    
    # Generate output filename
    base_name = os.path.splitext(filename)[0]
    output_filename = f"converted_{unique_id}_{base_name}.{target_format}"
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
    
    # Handle conversion
    from converter_universal import convert_image_to_pdf, convert_image_to_image
    
    if target_format == 'pdf':
        success, message = convert_image_to_pdf(input_path, output_path)
    else:
        success, message = convert_image_to_image(input_path, output_path, target_format, quality)
    
    if success and os.path.exists(output_path):
        return jsonify({
            'success': True,
            'message': message,
            'download_url': f'/download/{output_filename}',
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'error': message}), 500

@app.route('/translate', methods=['POST'])
def translate_file_route():
    """Translate uploaded documents with robust error handling"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        target_lang = request.form.get('target_lang', 'es')
        source_lang = request.form.get('source_lang', 'auto')
        
        # Validate file size (50MB max)
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > 50 * 1024 * 1024:
            return jsonify({'success': False, 'error': 'File too large (max 50MB)'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        unique_id = str(uuid.uuid4())
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(input_path)
        
        # Generate output path - SAFE INTERNAL NAME (avoid brackets for sanitization)
        file_ext = os.path.splitext(filename)[1].lower()
        base_name = os.path.splitext(filename)[0]
        
        # Internal name used for filesystem (safe)
        safe_base = secure_filename(base_name)
        
        # Standardize internal extensions for preview compatibility
        # Legacy .doc is bridged to .docx internally
        # Legacy .xls is bridged to .xlsx internally
        target_ext = file_ext
        if file_ext == '.doc':
            target_ext = '.docx'
        elif file_ext == '.xls':
            target_ext = '.xlsx'
            
        internal_filename = f"trans_{unique_id}_Translated_{safe_base}{target_ext}"
        if file_ext == '.pdf':
            internal_filename = f"trans_{unique_id}_Translated_{safe_base}.pdf"
            
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], internal_filename)
        
        # Pretty name used for user download
        download_name = f"[Translated]_{filename}"
        if file_ext == '.pdf':
            download_name = f"[Translated]_{base_name}.pdf"
        
        # Import translator function
        from translator_engine import translate_document
        
        # Call the translator function
        result = translate_document(
            input_path, output_path, target_lang, source_lang, file_ext
        )
        
        # After translation, if it succeeded, we want the client to download it with our clean name
        # download_name is already prepared above

        # Handle both 2-value and 3-value returns
        if isinstance(result, tuple) and len(result) == 2:
            success, message = result
            res_path = output_path if success else None
        elif isinstance(result, tuple) and len(result) == 3:
            success, message, res_path = result
        else:
            # Log unexpected return
            try:
                print(f"Unexpected return from translate_document: {result}")
            except UnicodeEncodeError:
                print("Unexpected return from translate_document: [Unicode Content]")
            return jsonify({'success': False, 'error': 'Internal server error'}), 500
        
        # Update output_filename for response
        output_filename = os.path.basename(output_path)
        
        if success and res_path and os.path.exists(res_path):
            # Verify the translation is valid for the target language
            try:
                with open(res_path, 'r', encoding='utf-8-sig', errors='ignore') as f:
                    translated_sample = f.read(1000)
                
                with open(input_path, 'r', encoding='utf-8-sig', errors='ignore') as f:
                    original_sample = f.read(1000) 
                
                from translator_engine import is_valid_translation
                if not is_valid_translation(translated_sample, target_lang, original_sample):
                    print(f"WARNING: Translation to {target_lang} may be invalid")
            except Exception as ve:
                pass
            
        # Updated logic: Use the exact path returned by the engine
        if success and res_path and os.path.exists(res_path):
            # The engine might have changed the extension (e.g., .pdf -> .txt)
            final_filename = os.path.basename(res_path)
            
            # Update download name if extension changed
            final_ext = os.path.splitext(final_filename)[1].lower()
            if final_ext != file_ext:
                download_name = os.path.splitext(download_name)[0] + final_ext
            
            # URL encode the download name for stability in headers
            encoded_download_name = urllib.parse.quote(download_name)
            
            return jsonify({
                'success': True,
                'message': message,
                'download_url': f'/download/{final_filename}?display_name={encoded_download_name}',
                'filename': download_name,
                'preview_filename': final_filename
            })
        else:
            return jsonify({'success': False, 'error': message or 'Translation failed'}), 500
            
    except Exception as e:
        try:
            print(f"Translation route error: {e}")
            import traceback
            traceback.print_exc()
        except UnicodeEncodeError:
            pass
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/get-preview/<filename>')
def get_preview(filename):
    """Retrieve text content for previewing"""
    filename = secure_filename(filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(file_path):
        return jsonify({'success': False, 'error': 'File not found'}), 404
        
    ext = os.path.splitext(filename)[1].lower().replace('.', '')
    
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8-sig', errors='ignore') as f:
                content = f.read(10000) # Preview first 10k chars
            return jsonify({'success': True, 'content': content, 'type': 'text'})
            
        elif ext in ['docx', 'doc']:
            from docx import Document
            
            # Legacy .doc preview bridge
            if ext == 'doc':
                try:
                    import pypandoc
                    temp_preview_docx = file_path + ".preview.docx"
                    pypandoc.convert_file(file_path, 'docx', outputfile=temp_preview_docx)
                    doc = Document(temp_preview_docx)
                    if os.path.exists(temp_preview_docx):
                        os.remove(temp_preview_docx)
                except Exception as pe:
                    return jsonify({'success': False, 'error': f'Legacy DOC preview conversion failed: {str(pe)}'})
            else:
                doc = Document(file_path)
                
            full_text = []
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract from tables (many DOCs have content here)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            content = "\n".join(full_text)[:10000]
            if not content.strip():
                content = "[Document structure found, but no text could be extracted for preview. Please download the file to view.]"
            return jsonify({'success': True, 'content': content, 'type': 'text'})
            
        elif ext == 'pdf':
            # Use pdfplumber for robust text extraction (handles non-Latin scripts like Hindi much better)
            import pdfplumber
            content = ""
            with pdfplumber.open(file_path) as pdf:
                # Get text from first few pages to show in preview
                for page in pdf.pages[:3]:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
            return jsonify({'success': True, 'content': content[:10000], 'type': 'text'})
            
        elif ext in ['xlsx', 'xls']:
            import pandas as pd
            # Use pandas for robust support of both .xls and .xlsx
            engine = 'xlrd' if ext == 'xls' else 'openpyxl'
            df = pd.read_excel(file_path, nrows=50, engine=engine)
            
            # Simple tab-separated text representation for preview
            content = df.to_csv(sep='\t', index=False)
            return jsonify({'success': True, 'content': content[:10000], 'type': 'text'})
            
        return jsonify({'success': False, 'error': f'Preview not available for {ext} format'})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Preview system error: {str(e)}'})

@app.route('/download/<filename>')
def download_file(filename):
    display_name = request.args.get('display_name')
    if display_name:
        # Unquote because it's coming from URL param
        display_name = urllib.parse.unquote(display_name)
        return send_from_directory(
            app.config['UPLOAD_FOLDER'], 
            filename, 
            as_attachment=True,
            download_name=display_name
        )
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
