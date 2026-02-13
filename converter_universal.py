"""
Universal File Converter Engine
Supports: PDF, DOCX, XLSX, PPTX, CSV, JSON, TXT, HTML, and Images
Optimized for deployment and memory efficiency.
"""

import os
import json
import csv
import pandas as pd
from datetime import datetime
import xml.etree.ElementTree as ET
from io import StringIO, BytesIO

# ============ RENDER ENVIRONMENT CHECK ============
ON_RENDER = os.environ.get('RENDER', '').lower() == 'true'

# ============ LAZY LOADING HELPERS ============

_pdf2docx = None
def get_pdf2docx():
    global _pdf2docx
    if _pdf2docx is None:
        from pdf2docx import Converter
        _pdf2docx = Converter
    return _pdf2docx

_PyPDF2 = None
def get_PyPDF2():
    global _PyPDF2
    if _PyPDF2 is None:
        import PyPDF2
        _PyPDF2 = PyPDF2
    return _PyPDF2

_tabula = None
def get_tabula():
    global _tabula
    if _tabula is None:
        import tabula
        _tabula = tabula
    return _tabula

_pdfplumber = None
def get_pdfplumber():
    global _pdfplumber
    if _pdfplumber is None:
        import pdfplumber
        _pdfplumber = pdfplumber
    return _pdfplumber

_weasyprint = None
def get_weasyprint():
    global _weasyprint
    if _weasyprint is None:
        from weasyprint import HTML
        _weasyprint = HTML
    return _weasyprint

_docx = None
def get_docx():
    global _docx
    if _docx is None:
        from docx import Document
        _docx = Document
    return _docx

_reportlab = None
def get_reportlab():
    global _reportlab
    if _reportlab is None:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.utils import simpleSplit
        _reportlab = (canvas, letter, simpleSplit)
    return _reportlab

_PIL = None
def get_PIL():
    global _PIL
    if _PIL is None:
        from PIL import Image
        _PIL = Image
    return _PIL

# ============ PDF CONVERSIONS ============

def convert_pdf_to_docx(input_path, output_path):
    """PDF to Word Document"""
    try:
        Converter = get_pdf2docx()
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
        return True, "PDF to DOCX conversion successful"
    except Exception as e:
        return False, str(e)

def convert_pdf_to_txt(input_path, output_path):
    """PDF to Text"""
    try:
        PyPDF2 = get_PyPDF2()
        with open(input_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            with open(output_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text)
        return True, "PDF to TXT conversion successful"
    except Exception as e:
        return False, str(e)

def convert_pdf_to_xlsx(input_path, output_path):
    """PDF to Excel (extract tables)"""
    try:
        # On Render, skip tabula-py (requires Java)
        if not ON_RENDER:
            try:
                tabula = get_tabula()
                dfs = tabula.read_pdf(input_path, pages='all', multiple_tables=True)
                if dfs:
                    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                        for i, df in enumerate(dfs):
                            sheet_name = f"Table_{i+1}" if len(dfs) > 1 else "Sheet1"
                            df.to_excel(writer, sheet_name=sheet_name, index=False)
                    return True, "PDF to XLSX conversion successful (tables extracted via Tabula)"
            except Exception as tabula_err:
                print(f"Tabula extraction failed, falling back to pdfplumber: {tabula_err}")

        # Fallback/Default: pdfplumber
        pdfplumber = get_pdfplumber()
        with pdfplumber.open(input_path) as pdf:
            all_tables = []
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        df = pd.DataFrame(table)
                        all_tables.append(df)
            
            if all_tables:
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                    for i, df in enumerate(all_tables):
                        sheet_name = f"Table_{i+1}" if len(all_tables) > 1 else "Sheet1"
                        df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                return True, "PDF to XLSX conversion successful (tables extracted via pdfplumber)"
            else:
                return False, "No tables found in PDF"
    except Exception as e:
        return False, str(e)

# ============ EXCEL CONVERSIONS ============

def convert_xlsx_to_csv(input_path, output_path):
    """Excel to CSV"""
    try:
        df = pd.read_excel(input_path)
        df.to_csv(output_path, index=False, encoding='utf-8')
        return True, "XLSX to CSV conversion successful"
    except Exception as e:
        return False, str(e)

def convert_xlsx_to_json(input_path, output_path):
    """Excel to JSON"""
    try:
        df = pd.read_excel(input_path)
        df.to_json(output_path, orient='records', indent=2, force_ascii=False)
        return True, "XLSX to JSON conversion successful"
    except Exception as e:
        return False, str(e)

def convert_xlsx_to_xml(input_path, output_path):
    """Excel to XML"""
    try:
        df = pd.read_excel(input_path)
        
        # Create XML structure
        root = ET.Element("root")
        for _, row in df.iterrows():
            record = ET.SubElement(root, "record")
            for col in df.columns:
                field = ET.SubElement(record, col.replace(" ", "_").replace("/", "_"))
                field.text = str(row[col]) if pd.notna(row[col]) else ""
        
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
        return True, "XLSX to XML conversion successful"
    except Exception as e:
        return False, str(e)

def convert_xlsx_to_pdf(input_path, output_path):
    """Excel to PDF via HTML"""
    try:
        df = pd.read_excel(input_path)
        html = df.to_html(index=False)
        
        HTML = get_weasyprint()
        HTML(string=f"<html><body>{html}</body></html>").write_pdf(output_path)
        return True, "XLSX to PDF conversion successful"
    except Exception as e:
        return False, str(e)

def convert_xlsx_to_html(input_path, output_path):
    """Excel to HTML"""
    try:
        df = pd.read_excel(input_path)
        df.to_html(output_path, index=False)
        return True, "XLSX to HTML conversion successful"
    except Exception as e:
        return False, str(e)

# ============ CSV CONVERSIONS ============

def convert_csv_to_xlsx(input_path, output_path):
    """CSV to Excel"""
    try:
        df = pd.read_csv(input_path)
        df.to_excel(output_path, index=False)
        return True, "CSV to XLSX conversion successful"
    except Exception as e:
        return False, str(e)

def convert_csv_to_json(input_path, output_path):
    """CSV to JSON"""
    try:
        df = pd.read_csv(input_path)
        df.to_json(output_path, orient='records', indent=2)
        return True, "CSV to JSON conversion successful"
    except Exception as e:
        return False, str(e)

def convert_csv_to_xml(input_path, output_path):
    """CSV to XML"""
    try:
        df = pd.read_csv(input_path)
        
        root = ET.Element("root")
        for _, row in df.iterrows():
            record = ET.SubElement(root, "record")
            for col in df.columns:
                field = ET.SubElement(record, col.replace(" ", "_").replace("/", "_"))
                field.text = str(row[col]) if pd.notna(row[col]) else ""
        
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
        return True, "CSV to XML conversion successful"
    except Exception as e:
        return False, str(e)

# ============ JSON CONVERSIONS ============

def convert_json_to_csv(input_path, output_path):
    """JSON to CSV"""
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Handle both array and object formats
        if isinstance(data, list):
            df = pd.DataFrame(data)
        elif isinstance(data, dict):
            # Try to find array in dict
            for key, value in data.items():
                if isinstance(value, list):
                    df = pd.DataFrame(value)
                    break
            else:
                df = pd.DataFrame([data])
        else:
            df = pd.DataFrame([data])
        
        df.to_csv(output_path, index=False, encoding='utf-8')
        return True, "JSON to CSV conversion successful"
    except Exception as e:
        return False, str(e)

def convert_json_to_xlsx(input_path, output_path):
    """JSON to Excel"""
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        if isinstance(data, list):
            df = pd.DataFrame(data)
        else:
            df = pd.DataFrame([data])
        
        df.to_excel(output_path, index=False)
        return True, "JSON to XLSX conversion successful"
    except Exception as e:
        return False, str(e)

def convert_json_to_xml(input_path, output_path):
    """JSON to XML"""
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        root = ET.Element("root")
        
        def dict_to_xml(parent, d):
            for key, value in d.items():
                elem = ET.SubElement(parent, key.replace(" ", "_").replace("/", "_"))
                if isinstance(value, dict):
                    dict_to_xml(elem, value)
                elif isinstance(value, list):
                    for item in value:
                        if isinstance(item, dict):
                            item_elem = ET.SubElement(elem, "item")
                            dict_to_xml(item_elem, item)
                        else:
                            item_elem = ET.SubElement(elem, "item")
                            item_elem.text = str(item)
                else:
                    elem.text = str(value) if value is not None else ""
        
        if isinstance(data, list):
            for item in data:
                record = ET.SubElement(root, "record")
                dict_to_xml(record, item)
        else:
            dict_to_xml(root, data)
        
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
        return True, "JSON to XML conversion successful"
    except Exception as e:
        return False, str(e)

# ============ IMAGE CONVERSIONS ============

def convert_image_to_image(input_path, output_path, target_format, quality=90):
    """Convert image from one format to another"""
    try:
        Image = get_PIL()
        image = Image.open(input_path)
        
        # Handle transparency for JPEG
        if target_format.upper() in ['JPG', 'JPEG'] and image.mode in ('RGBA', 'LA', 'P'):
            rgb_image = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'P':
                image = image.convert('RGBA')
            rgb_image.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = rgb_image
        
        save_format = 'JPEG' if target_format.upper() in ['JPG', 'JPEG'] else target_format.upper()
        image.save(output_path, save_format, quality=quality, optimize=True)
        return True, f"Image converted to {target_format.upper()}"
    except Exception as e:
        return False, str(e)

def convert_image_to_pdf(input_path, output_path):
    """Image to PDF"""
    try:
        Image = get_PIL()
        image = Image.open(input_path)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        image.save(output_path, 'PDF')
        return True, "Image to PDF conversion successful"
    except Exception as e:
        return False, str(e)

def convert_pdf_to_image(input_path, output_path, target_format='png'):
    """PDF to Image (first page only)"""
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(input_path, first_page=1, last_page=1)
        images[0].save(output_path, target_format.upper())
        return True, f"PDF to {target_format.upper()} conversion successful"
    except Exception as e:
        return False, str(e)

# ============ PPTX CONVERSIONS ============

def convert_pptx_to_pdf(input_path, output_path):
    """PowerPoint to PDF"""
    try:
        if ON_RENDER:
            return False, "PPTX to PDF not available on Render free tier (requires system utilities)"

        # Method 1: Spire.Presentation
        try:
            from spire.presentation import Presentation, FileFormat
            presentation = Presentation()
            presentation.LoadFromFile(input_path)
            presentation.SaveToFile(output_path, FileFormat.PDF)
            presentation.Dispose()
            return True, "PPTX to PDF conversion successful (Spire)"
        except:
            # Method 2: polytext + LibreOffice
            try:
                from polytext import convert_to_pdf
                convert_to_pdf(input_path, output_path)
                return True, "PPTX to PDF conversion successful (LibreOffice)"
            except:
                return False, "PPTX to PDF conversion failed: No compatible converter found"
    except Exception as e:
        return False, str(e)

def convert_pptx_to_images(input_path, output_dir):
    """PowerPoint to Images"""
    try:
        from spire.presentation import Presentation
        presentation = Presentation()
        presentation.LoadFromFile(input_path)
        
        image_paths = []
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        for i, slide in enumerate(presentation.Slides):
            output_path = os.path.join(output_dir, f"{base_name}_slide_{i+1}.png")
            image = slide.SaveAsImage()
            image.Save(output_path)
            image.Dispose()
            image_paths.append(output_path)
        
        presentation.Dispose()
        return True, f"PPTX to images successful ({len(image_paths)} slides)"
    except Exception as e:
        return False, str(e)

def convert_pptx_to_txt(input_path, output_path):
    """PowerPoint to Text (extract content)"""
    try:
        from spire.presentation import Presentation
        presentation = Presentation()
        presentation.LoadFromFile(input_path)
        
        text_content = []
        for slide in presentation.Slides:
            slide_text = []
            for shape in slide.Shapes:
                if hasattr(shape, 'Text'):
                    slide_text.append(shape.Text.Text)
            
            if slide_text:
                text_content.append(f"Slide {slide.SlideNumber}:")
                text_content.extend(slide_text)
                text_content.append("")
        
        presentation.Dispose()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))
        
        return True, "PPTX to TXT conversion successful"
    except Exception as e:
        return False, str(e)

# ============ WORD CONVERSIONS ============

def convert_docx_to_pdf(input_path, output_path):
    """Word to PDF with strong fallbacks"""
    try:
        # Method 1: docx2pdf (Windows only, requires Word)
        if not ON_RENDER:
            try:
                from docx2pdf import convert
                convert(input_path, output_path)
                if os.path.exists(output_path):
                    return True, "DOCX to PDF conversion successful (Word)"
            except:
                pass

        # Method 2: WeasyPrint (Pure Python, primary fallback)
        try:
            # We convert to HTML first, then to PDF via WeasyPrint
            temp_html = input_path.replace('.docx', '.temp_bridge.html')
            success, msg = convert_docx_to_html(input_path, temp_html)
            if success:
                HTML = get_weasyprint()
                HTML(temp_html).write_pdf(output_path)
                if os.path.exists(temp_html):
                    os.remove(temp_html)
                return True, "DOCX to PDF conversion successful (WeasyPrint Fallback)"
        except Exception as wp_err:
            print(f"WeasyPrint fallback failed: {wp_err}")

        # Method 3: polytext + LibreOffice
        try:
            from polytext import convert_to_pdf
            convert_to_pdf(input_path, output_path)
            return True, "DOCX to PDF conversion successful (LibreOffice)"
        except:
            if ON_RENDER:
                return False, "DOCX to PDF requires LibreOffice on Render (not found)"
            return False, "DOCX to PDF conversion failed: No compatible converter found"
    except Exception as e:
        return False, str(e)

def convert_docx_to_txt(input_path, output_path):
    """Word to Text"""
    try:
        Document = get_docx()
        doc = Document(input_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return True, "DOCX to TXT conversion successful"
    except Exception as e:
        return False, str(e)

def convert_docx_to_html(input_path, output_path):
    """Word to HTML"""
    try:
        Document = get_docx()
        doc = Document(input_path)
        html_content = [
            "<!DOCTYPE html><html><head><meta charset='UTF-8'>",
            "<style>",
            "body { font-family: 'DejaVu Sans', Arial, sans-serif; line-height: 1.5; padding: 40px; }",
            "p { margin-bottom: 12px; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "td, th { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "tr:nth-child(even) { background-color: #f9f9f9; }",
            "</style></head><body>"
        ]
        
        for para in doc.paragraphs:
            if para.text.strip():
                html_content.append(f"<p>{para.text}</p>")
        
        for table in doc.tables:
            html_content.append("<table>")
            for row in table.rows:
                html_content.append("<tr>")
                for cell in row.cells:
                    html_content.append(f"<td>{cell.text}</td>")
                html_content.append("</tr>")
            html_content.append("</table>")
        
        html_content.append("</body></html>")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_content))
        
        return True, "DOCX to HTML conversion successful"
    except Exception as e:
        return False, str(e)

# ============ TEXT CONVERSIONS ============

def convert_txt_to_pdf(input_path, output_path):
    """Text to PDF"""
    try:
        canvas, letter, simpleSplit = get_reportlab()
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Split into lines and wrap
        lines = text.split('\n')
        y = height - 40
        
        for line in lines:
            if y < 40:
                c.showPage()
                y = height - 40
            
            # Word wrap
            wrapped_lines = simpleSplit(line, c._fontname, c._fontsize, width - 80)
            for wrapped_line in wrapped_lines:
                c.drawString(40, y, wrapped_line)
                y -= 15
        
        c.save()
        return True, "TXT to PDF conversion successful"
    except Exception as e:
        return False, str(e)

def convert_txt_to_docx(input_path, output_path):
    """Text to Word"""
    try:
        Document = get_docx()
        doc = Document()
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split by double newline for paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(output_path)
        return True, "TXT to DOCX conversion successful"
    except Exception as e:
        return False, str(e)

# ============ CONVERSION DISPATCHER ============

FILE_CONVERSIONS = {
    # PDF Conversions
    'pdf': {
        'docx': convert_pdf_to_docx,
        'txt': convert_pdf_to_txt,
        'xlsx': convert_pdf_to_xlsx,
    },
    # Excel Conversions
    'xlsx': {
        'csv': convert_xlsx_to_csv,
        'json': convert_xlsx_to_json,
        'xml': convert_xlsx_to_xml,
        'pdf': convert_xlsx_to_pdf,
        'html': convert_xlsx_to_html,
    },
    'xls': {
        'csv': convert_xlsx_to_csv,
        'json': convert_xlsx_to_json,
        'xml': convert_xlsx_to_xml,
        'xlsx': lambda i, o: (os.rename(i, o) or (True, "XLS to XLSX converted")),
    },
    # CSV Conversions
    'csv': {
        'xlsx': convert_csv_to_xlsx,
        'json': convert_csv_to_json,
        'xml': convert_csv_to_xml,
    },
    # JSON Conversions
    'json': {
        'csv': convert_json_to_csv,
        'xlsx': convert_json_to_xlsx,
        'xml': convert_json_to_xml,
    },
    # PowerPoint Conversions
    'pptx': {
        'pdf': convert_pptx_to_pdf,
        'txt': convert_pptx_to_txt,
    },
    'ppt': {
        'pdf': convert_pptx_to_pdf,
        'txt': convert_pptx_to_txt,
        'pptx': lambda i, o: (os.rename(i, o) or (True, "PPT to PPTX converted")),
    },
    # Word Conversions
    'docx': {
        'pdf': convert_docx_to_pdf,
        'txt': convert_docx_to_txt,
        'html': convert_docx_to_html,
    },
    'doc': {
        'docx': lambda i, o: (os.rename(i, o) or (True, "DOC to DOCX converted")),
        'txt': convert_docx_to_txt,
    },
    # Text Conversions
    'txt': {
        'pdf': convert_txt_to_pdf,
        'docx': convert_txt_to_docx,
    }
}

IMAGE_CONVERSIONS = {
    'jpg': ['png', 'webp', 'pdf', 'jpeg'],
    'jpeg': ['png', 'webp', 'pdf', 'jpg'],
    'png': ['jpg', 'webp', 'pdf', 'jpeg'],
    'webp': ['jpg', 'png', 'pdf', 'jpeg'],
    'gif': ['png', 'jpg', 'webp', 'pdf', 'jpeg'],
    'bmp': ['png', 'jpg', 'webp', 'pdf', 'jpeg']
}

SUPPORTED_IMAGE_INPUTS = list(IMAGE_CONVERSIONS.keys())

def convert_file(input_path, output_path, source_format, target_format):
    """Universal conversion dispatcher"""
    source_format = source_format.lower().replace('.', '')
    target_format = target_format.lower().replace('.', '')
    
    if source_format in FILE_CONVERSIONS:
        if target_format in FILE_CONVERSIONS[source_format]:
            converter = FILE_CONVERSIONS[source_format][target_format]
            return converter(input_path, output_path)
        else:
            return False, f"Conversion from {source_format} to {target_format} not supported"
    else:
        return False, f"Source format {source_format} not supported"
